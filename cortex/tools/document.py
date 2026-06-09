"""
cortex.tools.document
──────────────────────
Create formatted documents: .docx (Word) and .txt.

Requires python-docx for Word files:
    pip install python-docx

The model passes structured content; this tool handles formatting.
"""

from __future__ import annotations

import os
from pathlib import Path

SCHEMA = {
    "type": "function",
    "function": {
        "name": "document",
        "description": (
            "Create a formatted Word (.docx) or plain text (.txt) document. "
            "Use this instead of filesystem(write) when the user asks for a Word document, "
            "a .docx file, or a document with headings/formatting. "
            "Pass the full content as markdown-like text — headings with #/## become "
            "Word headings, **bold** becomes bold, bullet lines become lists."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Output file path. Use .docx for Word, .txt for plain text. "
                                   "Example: ~/Desktop/report.docx",
                },
                "title": {
                    "type": "string",
                    "description": "Document title (shown as the first heading).",
                },
                "content": {
                    "type": "string",
                    "description": (
                        "Document body. Supports:\n"
                        "  # Heading 1\n"
                        "  ## Heading 2\n"
                        "  ### Heading 3\n"
                        "  - bullet item\n"
                        "  **bold text**\n"
                        "  regular paragraph text"
                    ),
                },
            },
            "required": ["path", "content"],
        },
    },
}


def _resolve(path: str) -> Path:
    return Path(os.path.expandvars(path)).expanduser().resolve()


def _write_txt(p: Path, title: str, content: str) -> str:
    """Fallback: write clean UTF-8 .txt."""
    lines = []
    if title:
        lines += [title, "=" * len(title), ""]
    lines.append(content)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text("\n".join(lines), encoding="utf-8")
    return f"Created {p}  ({p.stat().st_size // 1024 + 1} KB)"


def _write_docx(p: Path, title: str, content: str) -> str:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # Graceful fallback to .txt with a clear message
        txt_path = p.with_suffix(".txt")
        result = _write_txt(txt_path, title, content)
        return (
            f"[NOTE] python-docx not installed — created plain text instead.\n"
            f"Install: pip install python-docx\n{result}"
        )

    doc = Document()

    # Title
    if title:
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse content line by line
    for line in content.splitlines():
        stripped = line.strip()

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p_obj = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped == "":
            doc.add_paragraph("")
        else:
            # Handle inline **bold**
            para = doc.add_paragraph()
            _add_inline(para, stripped)

    p.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(p))
    return f"Created {p}  ({p.stat().st_size // 1024 + 1} KB)"


def _add_inline(para, text: str) -> None:
    """Add a paragraph with **bold** support."""
    import re
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def execute(path: str, content: str, title: str = "", **_) -> str:
    try:
        p = _resolve(path)
        suffix = p.suffix.lower()

        if suffix == ".docx":
            return _write_docx(p, title, content)
        else:
            # .txt or anything else — always UTF-8
            return _write_txt(p, title, content)

    except Exception as e:
        return f"[ERROR] document: {type(e).__name__}: {e}"
